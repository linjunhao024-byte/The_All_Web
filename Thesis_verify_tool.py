"""
毕业论文排版自动校验引擎

功能：读取 .docx 文件，执行以下校验并生成结构化报告：
  1. 全局页边距（上/下 2.54cm，左/右 3.17cm）
  2. 文档结构顺序（封面→声明→目录→摘要→正文→参考文献→致谢→附录）
  3. 封面格式（主标题宋体/一号/加粗，日期阿拉伯数字）
  4. 中文题目与摘要（楷体_GB2312、二号/小四、居中）
  5. 英文题目与摘要（Times New Roman、二号/加粗/居中）
  6. 正文字体与行距（宋体/TNR、小四、1.5 倍行距）
  7. 多级标题样式与缩进（黑体、字号、双轨制缩进规则）
  8. 图表题注位置与格式（表上方/图下方、居中、五号）
  9. 参考文献比例指标（总数/近3年/外文/学位论文）
  10. 脚注格式（宋体、小五号）

依赖：python-docx (外部), re/sys/datetime (标准库)
用法：python Thesis_verify_tool.py <path_to_docx>
"""

import re
import sys
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, field

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn


# ─────────────────────────────────────────────────────────────
# 数据结构
# ─────────────────────────────────────────────────────────────

@dataclass
class CheckResult:
    """单条校验结果"""
    name: str
    passed: bool
    message: str
    context_text: Optional[str] = None


@dataclass
class AnalysisReport:
    """最终校验报告"""
    passed_items: List[Dict[str, str]] = field(default_factory=list)
    failed_items: List[Dict[str, str]] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)

    def add_pass(self, name: str, message: str):
        self.passed_items.append({"name": name, "message": message})

    def add_fail(self, name: str, message: str, context_text: str = None):
        item = {"name": name, "message": message}
        if context_text:
            item["context_text"] = context_text
        self.failed_items.append(item)

    def add_error(self, message: str):
        self.errors.append(message)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "passed_items": self.passed_items,
            "failed_items": self.failed_items,
            "errors": self.errors,
            "summary": {
                "total_checks": len(self.passed_items) + len(self.failed_items),
                "passed": len(self.passed_items),
                "failed": len(self.failed_items),
                "errors": len(self.errors),
            },
        }


# ─────────────────────────────────────────────────────────────
# ThesisFormatVerifier —— 毕业论文格式校验引擎
# ─────────────────────────────────────────────────────────────

class ThesisFormatVerifier:
    """
    毕业论文排版校验器。

    阶段一职责：
    - 校验全局页边距
    - 校验文档结构化排序
    """

    # 文档结构关键词锚点（按正确顺序排列）
    # 使用列表保持顺序，每个元素为 (章节名, 关键词列表)
    STRUCTURE_ANCHORS = [
        ("封面", ["封面", "毕业论文", "毕业设计"]),
        ("声明", ["声明", "郑重声明", "独创性声明", "原创性"]),
        ("目录", ["目录", "目 录"]),
        ("中英文摘要", ["摘要", "摘 要", "Abstract", "ABSTRACT"]),
        ("正文", ["正文", "第一章", "第二章", "第1章", "第2章",
                 "1 ", "1.", "绪论", "引言"]),
        ("参考文献", ["参考文献"]),
        ("致谢", ["致谢", "致 谢", "答谢", "谢辞"]),
        ("附录", ["附录", "附 录"]),
    ]

    # 页边距标准值（单位：厘米）
    MARGIN_TOP = 2.54
    MARGIN_BOTTOM = 2.54
    MARGIN_LEFT = 3.17
    MARGIN_RIGHT = 3.17
    # 浮点容差（0.1 毫米 = 0.01 厘米）
    MARGIN_TOLERANCE = 0.01

    # ── 字号标准值（单位：磅 pt）──
    FONT_SIZE_MAP = {
        "一号": 36.0, "小一": 24.0, "二号": 22.0, "小二": 18.0,
        "三号": 16.0, "小三": 15.0, "四号": 14.0, "小四": 12.0,
        "五号": 10.5, "小五": 9.0,
    }
    FONT_SIZE_TOLERANCE = 0.5  # 字号容差 ±0.5pt

    # ── 标题序号正则 ──
    # 阿拉伯数字格式：1、1.1、1.1.1、1.1.1.1
    RE_HEADING_ARABIC = re.compile(
        r"^\s*\d+(\.\d+){0,3}\s*[、..,、]?\s*\S"
    )
    # 中文序号格式：一、（一）、（1）、①
    RE_HEADING_CHINESE = re.compile(
        r"^\s*[一二三四五六七八九十百]+[、．.]\s*\S"
        r"|^\s*（[一二三四五六七八九十百]+）\s*\S"
        r"|^\s*\([一二三四五六七八九十百]+\)\s*\S"
        r"|^\s*（\d+）\s*\S"
        r"|^\s*\(\d+\)\s*\S"
        r"|^\s*[①②③④⑤⑥⑦⑧⑨⑩]\s*\S"
    )
    # 三级标题典型序号：1.1.1 或（1）
    RE_HEADING_LEVEL3 = re.compile(
        r"^\s*\d+\.\d+\.\d+\s*"
        r"|^\s*（\d+）\s*"
        r"|^\s*\(\d+\)\s*"
    )

    def __init__(self, doc_path: str):
        """
        初始化校验器，加载 .docx 文档。

        Args:
            doc_path: .docx 文件路径
        """
        self.doc_path = doc_path
        self.doc: Optional[Document] = None
        self.report = AnalysisReport()
        self._load_document()

    def _load_document(self):
        """加载文档，捕获异常"""
        try:
            self.doc = Document(self.doc_path)
        except Exception as e:
            self.report.add_error(f"文件读取失败: {str(e)}")

    def run_all(self) -> Dict[str, Any]:
        """
        执行全部校验项，返回统一格式的报告字典。

        Returns:
            {
                "passed_items": [{"name": "...", "message": "..."}, ...],
                "failed_items": [{"name": "...", "message": "..."}, ...],
                "errors": ["..."],
                "summary": {"total_checks": N, "passed": N, "failed": N, "errors": N}
            }
        """
        if self.doc is None:
            return self.report.to_dict()

        self._check_global_margins()
        self._check_document_structure()
        self._check_cover()
        self._check_chinese_title_abstract()
        self._check_english_title_abstract()
        self._check_fonts_and_spacing()
        self._check_heading_styles()
        self._check_tables_and_figures()
        self._check_references_metrics()
        self._check_footnotes()

        return self.report.to_dict()

    # ─────────────────────────────────────────────────────────
    # 校验项：全局页边距
    # ─────────────────────────────────────────────────────────

    def _check_global_margins(self):
        """
        校验全局页边距是否符合规范。

        规范：
        - 上 2.54 厘米，下 2.54 厘米，左 3.17 厘米，右 3.17 厘米
        - 允许 0.1 毫米（0.01 厘米）的浮点误差
        """
        sections = self.doc.sections
        if not sections:
            self.report.add_error("文档中未找到节(Section)，无法校验页边距")
            return

        # 遍历所有节，收集不合规的节
        non_compliant_sections = []
        for idx, section in enumerate(sections):
            issues = []

            # python-docx 中页边距单位为 EMU（English Metric Units）
            # 1 厘米 = 360000 EMU
            # 但 python-docx 的 Cm 类可以直接比较
            top_cm = section.top_margin / 360000 if section.top_margin is not None else None
            bottom_cm = section.bottom_margin / 360000 if section.bottom_margin is not None else None
            left_cm = section.left_margin / 360000 if section.left_margin is not None else None
            right_cm = section.right_margin / 360000 if section.right_margin is not None else None

            if top_cm is None:
                issues.append("上边距未设定")
            elif abs(top_cm - self.MARGIN_TOP) > self.MARGIN_TOLERANCE:
                issues.append(f"上边距应为 {self.MARGIN_TOP}cm，实际为 {top_cm:.2f}cm")

            if bottom_cm is None:
                issues.append("下边距未设定")
            elif abs(bottom_cm - self.MARGIN_BOTTOM) > self.MARGIN_TOLERANCE:
                issues.append(f"下边距应为 {self.MARGIN_BOTTOM}cm，实际为 {bottom_cm:.2f}cm")

            if left_cm is None:
                issues.append("左边距未设定")
            elif abs(left_cm - self.MARGIN_LEFT) > self.MARGIN_TOLERANCE:
                issues.append(f"左边距应为 {self.MARGIN_LEFT}cm，实际为 {left_cm:.2f}cm")

            if right_cm is None:
                issues.append("右边距未设定")
            elif abs(right_cm - self.MARGIN_RIGHT) > self.MARGIN_TOLERANCE:
                issues.append(f"右边距应为 {self.MARGIN_RIGHT}cm，实际为 {right_cm:.2f}cm")

            if issues:
                non_compliant_sections.append((idx + 1, issues))

        if not non_compliant_sections:
            self.report.add_pass(
                "全局页边距",
                f"全部 {len(sections)} 节的页边距均符合规范"
                f"（上/下 {self.MARGIN_TOP}cm，左/右 {self.MARGIN_LEFT}cm）"
            )
        else:
            detail_lines = []
            for sec_num, issues in non_compliant_sections:
                detail_lines.append(f"第 {sec_num} 节：" + "；".join(issues))
            self.report.add_fail(
                "全局页边距",
                f"共 {len(non_compliant_sections)} 节页边距不合规",
                context_text="\n".join(detail_lines)
            )

    # ─────────────────────────────────────────────────────────
    # 校验项：文档结构化排序
    # ─────────────────────────────────────────────────────────

    def _check_document_structure(self):
        """
        通过关键词锚定，粗略检查文档是否包含了 8 个部分的关键字，
        并且顺序正确。

        结构顺序：封面 → 声明 → 目录 → 中英文摘要 → 正文 → 参考文献 → 致谢 → 附录
        """
        if not self.doc.paragraphs:
            self.report.add_error("文档中未找到任何段落，无法校验结构顺序")
            return

        # 收集全文段落文本（保留顺序）
        all_text = [p.text.strip() for p in self.doc.paragraphs if p.text.strip()]

        if not all_text:
            self.report.add_error("文档段落均为空，无法校验结构顺序")
            return

        # 对每个结构锚点，找到其关键词首次出现的段落索引
        found_positions: List[tuple] = []  # (章节名, 最小索引, 匹配到的关键词)
        missing_sections: List[str] = []

        for section_name, keywords in self.STRUCTURE_ANCHORS:
            earliest_idx = len(all_text)  # 初始化为最大值
            matched_kw = None

            for kw in keywords:
                for i, text in enumerate(all_text):
                    if kw in text and i < earliest_idx:
                        earliest_idx = i
                        matched_kw = kw
                        # 不 break，因为要找最早的出现位置

            if earliest_idx < len(all_text):
                found_positions.append((section_name, earliest_idx, matched_kw))
            else:
                missing_sections.append(section_name)

        # 报告缺失的章节
        if missing_sections:
            self.report.add_fail(
                "文档结构完整性",
                f"缺少以下章节: {', '.join(missing_sections)}",
                context_text="文档应包含以下 8 个部分（按顺序）：\n"
                             + " → ".join(name for name, _ in self.STRUCTURE_ANCHORS)
            )
        else:
            self.report.add_pass(
                "文档结构完整性",
                "全部 8 个结构部分均已找到"
            )

        # 检查顺序：已找到的章节应按索引升序排列
        if len(found_positions) >= 2:
            out_of_order = []
            for i in range(1, len(found_positions)):
                prev_name, prev_idx, _ = found_positions[i - 1]
                curr_name, curr_idx, _ = found_positions[i]
                if curr_idx < prev_idx:
                    out_of_order.append(
                        f"「{curr_name}」(段落{curr_idx}) 出现在「{prev_name}」(段落{prev_idx}) 之前"
                    )

            if out_of_order:
                # 构建实际顺序上下文
                order_lines = ["文档中检测到的实际顺序："]
                for name, idx, kw in found_positions:
                    order_lines.append(f"  {name} — 关键字「{kw}」首次出现于段落 {idx}")

                self.report.add_fail(
                    "文档结构顺序",
                    f"存在 {len(out_of_order)} 处顺序错误",
                    context_text="\n".join(order_lines)
                )
            else:
                order_summary = " → ".join(name for name, _, _ in found_positions)
                self.report.add_pass(
                    "文档结构顺序",
                    f"结构顺序正确: {order_summary}"
                )

    # ─────────────────────────────────────────────────────────
    # 校验项：封面格式
    # ─────────────────────────────────────────────────────────

    def _check_cover(self):
        """
        校验封面格式：
        - 主标题必须包含"毕业论文（设计）"字样，宋体、一号(36pt)、加粗。
        - 日期栏必须使用阿拉伯数字格式（如 2025 年 11 月 8 日）。
        - 跳过"原创性声明及版权使用授权书"页（不校验）。
        """
        paras = self.doc.paragraphs
        if not paras:
            self.report.add_error("文档无段落，跳过封面校验")
            return

        # 定位封面区域：从文档开头到第一个正文锚点之前
        cover_end = len(paras)
        for i, p in enumerate(paras):
            text = p.text.strip()
            if any(kw in text for kw in ["第一章", "第1章", "绪论", "引言",
                                          "第 一 章", "目 录", "目录"]):
                cover_end = i
                break

        cover_paras = paras[:cover_end]
        if not cover_paras:
            self.report.add_error("无法定位封面区域，跳过封面校验")
            return

        # ── 跳过声明页 ──
        # 声明页关键字：原创性声明、版权使用授权书
        # 我们只在封面区域找主标题和日期，声明页内容会被自然跳过

        # ── 主标题校验 ──
        title_para = None
        for p in cover_paras:
            text = p.text.strip()
            clean = text.replace("（", "(").replace("）", ")").replace(" ", "")
            if "毕业论文" in clean or "毕业设计" in clean:
                # 排除声明页中的"毕业论文"字样
                if "原创性" not in text and "声明" not in text and "授权" not in text:
                    title_para = p
                    break

        if title_para is None:
            self.report.add_fail(
                "封面主标题",
                "未找到包含「毕业论文（设计）」的主标题段落"
            )
        else:
            font_name = None
            font_size = None
            bold = None
            for run in title_para.runs:
                if not run.text.strip():
                    continue
                if font_name is None:
                    ea = self._get_run_east_asian_font(run)
                    font_name = ea if ea else run.font.name
                if font_size is None and run.font.size:
                    font_size = run.font.size.pt
                if bold is None and run.bold is not None:
                    bold = run.bold

            issues = []
            if font_name and "宋体" not in font_name and "SimSun" not in font_name:
                issues.append(f"字体应为宋体，实际为 {font_name}")
            if font_size and abs(font_size - 36.0) > 1.0:
                issues.append(f"字号应为一号(36pt)，实际为 {font_size}pt")
            if bold is not None and not bold:
                issues.append("应为加粗，实际为非加粗")

            if not issues:
                self.report.add_pass(
                    "封面主标题",
                    f"主标题格式符合要求（宋体、一号、加粗）"
                )
            else:
                self.report.add_fail(
                    "封面主标题",
                    "；".join(issues)
                )

        # ── 日期栏校验 ──
        # 查找包含"年"和"月"和"日"的段落，检查是否使用阿拉伯数字
        date_para = None
        for p in cover_paras:
            text = p.text.strip()
            if "年" in text and "月" in text and "日" in text:
                if "原创性" not in text and "声明" not in text and "授权" not in text:
                    date_para = p
                    break

        if date_para is None:
            self.report.add_fail(
                "封面日期",
                "封面中未找到包含「年 月 日」的日期段落"
            )
        else:
            text = date_para.text.strip()
            # 检查是否使用阿拉伯数字：提取年月日前的数字部分
            date_match = re.search(r"(\d+)\s*年\s*(\d+)\s*月\s*(\d+)\s*日", text)
            if date_match:
                self.report.add_pass(
                    "封面日期",
                    f"日期使用阿拉伯数字格式: {date_match.group(0)}"
                )
            else:
                # 可能使用了中文数字（如"二〇二五年"）
                self.report.add_fail(
                    "封面日期",
                    f"日期应使用阿拉伯数字格式（如 2025 年 11 月 8 日），"
                    f"当前为「{text[:40]}」"
                )

    # ─────────────────────────────────────────────────────────
    # 校验项：中文题目与摘要页
    # ─────────────────────────────────────────────────────────

    def _check_chinese_title_abstract(self):
        """
        校验中文题目与摘要页：
        - 论文主标题：楷体_GB2312，二号(22pt)，居中。
        - 作者姓名：楷体_GB2312，小四号(12pt)。
        """
        paras = self.doc.paragraphs

        # 定位"摘要"区域
        abstract_start = None
        abstract_end = None
        for i, p in enumerate(paras):
            text = p.text.strip()
            if text in ("摘要", "摘 要", "摘　要"):
                abstract_start = i
                continue
            if abstract_start is not None:
                # 遇到 Abstract 或正文则停止
                if text in ("Abstract", "ABSTRACT", "第一章", "第1章", "绪论"):
                    abstract_end = i
                    break

        if abstract_start is None:
            self.report.add_error("未找到中文「摘要」段落，跳过中文题目与摘要校验")
            return

        if abstract_end is None:
            abstract_end = min(abstract_start + 30, len(paras))

        section_paras = paras[abstract_start:abstract_end]

        # ── 论文主标题校验 ──
        # 主标题通常是摘要标题之后、作者姓名之前的第一段非空文本
        title_para = None
        author_para = None
        found_abstract_title = False

        for p in section_paras:
            text = p.text.strip()
            if text in ("摘要", "摘 要", "摘　要"):
                found_abstract_title = True
                continue
            if found_abstract_title and text:
                if title_para is None:
                    title_para = p
                elif author_para is None:
                    # 第二段非空文本可能是作者姓名
                    # 作者姓名通常较短且不含标点
                    if len(text) < 30 and "。" not in text and "，" not in text:
                        author_para = p
                    break

        if title_para is None:
            self.report.add_fail(
                "中文题目",
                "摘要区域中未找到论文主标题段落"
            )
        else:
            font_name = None
            font_size = None
            alignment = title_para.alignment
            for run in title_para.runs:
                if not run.text.strip():
                    continue
                if font_name is None:
                    ea = self._get_run_east_asian_font(run)
                    font_name = ea if ea else run.font.name
                if font_size is None and run.font.size:
                    font_size = run.font.size.pt

            issues = []
            if font_name:
                if "楷体" not in font_name and "KaiTi" not in font_name:
                    issues.append(f"字体应为楷体_GB2312，实际为 {font_name}")
            if font_size and abs(font_size - 22.0) > 1.0:
                issues.append(f"字号应为二号(22pt)，实际为 {font_size}pt")
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            if alignment is not None and alignment != WD_ALIGN_PARAGRAPH.CENTER:
                issues.append("应居中对齐")

            if not issues:
                self.report.add_pass(
                    "中文题目",
                    f"论文主标题格式符合要求（楷体_GB2312、二号、居中）"
                )
            else:
                self.report.add_fail(
                    "中文题目",
                    "；".join(issues)
                )

        # ── 作者姓名校验 ──
        if author_para is None:
            self.report.add_fail(
                "中文摘要作者",
                "摘要区域中未找到作者姓名段落"
            )
        else:
            font_name = None
            font_size = None
            for run in author_para.runs:
                if not run.text.strip():
                    continue
                if font_name is None:
                    ea = self._get_run_east_asian_font(run)
                    font_name = ea if ea else run.font.name
                if font_size is None and run.font.size:
                    font_size = run.font.size.pt

            issues = []
            if font_name:
                if "楷体" not in font_name and "KaiTi" not in font_name:
                    issues.append(f"字体应为楷体_GB2312，实际为 {font_name}")
            if font_size and abs(font_size - 12.0) > 0.5:
                issues.append(f"字号应为小四(12pt)，实际为 {font_size}pt")

            if not issues:
                self.report.add_pass(
                    "中文摘要作者",
                    f"作者姓名格式符合要求（楷体_GB2312、小四）"
                )
            else:
                self.report.add_fail(
                    "中文摘要作者",
                    "；".join(issues)
                )

    # ─────────────────────────────────────────────────────────
    # 校验项：英文题目与摘要页
    # ─────────────────────────────────────────────────────────

    def _check_english_title_abstract(self):
        """
        校验英文题目与摘要页：
        - English Title：Times New Roman，二号(22pt)，加粗，居中。
        - 作者姓名（拼音）和日期：Times New Roman。
        """
        paras = self.doc.paragraphs

        # 定位 Abstract 区域
        abstract_start = None
        abstract_end = None
        for i, p in enumerate(paras):
            text = p.text.strip()
            if text in ("Abstract", "ABSTRACT"):
                abstract_start = i
                continue
            if abstract_start is not None:
                if text in ("Keywords", "KEYWORDS", "关键词", "Key words",
                            "第一章", "第1章", "绪论", "参考文献"):
                    abstract_end = i
                    break

        if abstract_start is None:
            self.report.add_error("未找到英文「Abstract」段落，跳过英文题目与摘要校验")
            return

        if abstract_end is None:
            abstract_end = min(abstract_start + 30, len(paras))

        section_paras = paras[abstract_start:abstract_end]

        # ── English Title 校验 ──
        # Abstract 标题之后的第一段非空文本通常是英文论文标题
        title_para = None
        found_abstract = False
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        for p in section_paras:
            text = p.text.strip()
            if text in ("Abstract", "ABSTRACT"):
                found_abstract = True
                continue
            if found_abstract and text:
                title_para = p
                break

        if title_para is None:
            self.report.add_fail(
                "英文题目",
                "Abstract 区域中未找到英文论文标题段落"
            )
        else:
            font_name = None
            font_size = None
            bold = None
            alignment = title_para.alignment
            for run in title_para.runs:
                if not run.text.strip():
                    continue
                if font_name is None:
                    font_name = run.font.name
                if font_size is None and run.font.size:
                    font_size = run.font.size.pt
                if bold is None and run.bold is not None:
                    bold = run.bold

            issues = []
            if font_name and "Times New Roman" not in font_name:
                issues.append(f"字体应为 Times New Roman，实际为 {font_name}")
            if font_size and abs(font_size - 22.0) > 1.0:
                issues.append(f"字号应为二号(22pt)，实际为 {font_size}pt")
            if bold is not None and not bold:
                issues.append("应为加粗，实际为非加粗")
            if alignment is not None and alignment != WD_ALIGN_PARAGRAPH.CENTER:
                issues.append("应居中对齐")

            if not issues:
                self.report.add_pass(
                    "英文题目",
                    f"English Title 格式符合要求（Times New Roman、二号、加粗、居中）"
                )
            else:
                self.report.add_fail(
                    "英文题目",
                    "；".join(issues)
                )

        # ── 作者姓名与日期校验 ──
        # 在 Abstract 之后，Keywords 之前，查找作者/日期段落
        author_date_paras = []
        for p in section_paras:
            text = p.text.strip()
            if text in ("Abstract", "ABSTRACT", "Keywords", "KEYWORDS", "关键词"):
                continue
            if text and p is not title_para:
                author_date_paras.append(p)

        if not author_date_paras:
            self.report.add_error("Abstract 区域中未找到作者/日期段落，跳过校验")
            return

        font_issues = []
        for p in author_date_paras:
            text = p.text.strip()
            for run in p.runs:
                if not run.text.strip():
                    continue
                font_name = run.font.name
                if font_name and "Times New Roman" not in font_name:
                    font_issues.append(
                        f"「{text[:30]}」应为 Times New Roman，实际为 {font_name}"
                    )
                    break  # 每段只报一次

        if not font_issues:
            self.report.add_pass(
                "英文摘要作者/日期",
                "作者姓名与日期字体均为 Times New Roman"
            )
        else:
            self.report.add_fail(
                "英文摘要作者/日期",
                f"共 {len(font_issues)} 处字体不合规",
                context_text="\n".join(font_issues[:10])
            )

    # ─────────────────────────────────────────────────────────
    # 校验项：脚注（注释）
    # ─────────────────────────────────────────────────────────

    def _check_footnotes(self):
        """
        校验脚注格式：
        - 如果文档中存在脚注，检查其字体是否为宋体，字号是否为小五号(9pt)。
        """
        # python-docx 没有直接遍历脚注的 API，
        # 需要通过 XML 访问 footnotes 部件
        try:
            footnotes_part = None
            for rel in self.doc.part.rels.values():
                if "footnotes" in rel.reltype:
                    footnotes_part = rel.target_part
                    break

            if footnotes_part is None:
                self.report.add_error("文档中未检测到脚注部分，跳过脚注校验")
                return

            # 解析脚注 XML
            from lxml import etree
            nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            root = etree.fromstring(footnotes_part.blob)
            footnote_elements = root.findall('.//w:footnote', nsmap)

            if not footnote_elements:
                self.report.add_error("脚注部分为空，跳过脚注校验")
                return

            font_issues = []
            scanned = 0

            for fn_elem in footnote_elements:
                # 跳过分隔线脚注（type="separator"）
                fn_type = fn_elem.get(qn('w:type'), '')
                if fn_type in ('separator', 'continuationSeparator'):
                    continue

                # 提取脚注文本段落
                para_elems = fn_elem.findall('.//w:p', nsmap)
                for para_elem in para_elems:
                    # 提取文本
                    texts = []
                    for t in para_elem.iter(qn('w:t')):
                        if t.text:
                            texts.append(t.text)
                    fn_text = "".join(texts).strip()
                    if not fn_text:
                        continue

                    scanned += 1
                    if scanned > 50:  # 安全上限
                        break

                    # 检查每个 run 的字体和字号
                    run_elems = para_elem.findall('.//w:r', nsmap)
                    for r_elem in run_elems:
                        run_text_parts = []
                        for t in r_elem.iter(qn('w:t')):
                            if t.text:
                                run_text_parts.append(t.text)
                        run_text = "".join(run_text_parts).strip()
                        if not run_text:
                            continue

                        # 获取字体
                        rPr = r_elem.find(qn('w:rPr'))
                        font_name = None
                        font_size = None

                        if rPr is not None:
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is not None:
                                ea = rFonts.get(qn('w:eastAsia'))
                                if ea:
                                    font_name = ea
                                elif rFonts.get(qn('w:ascii')):
                                    font_name = rFonts.get(qn('w:ascii'))

                            sz = rPr.find(qn('w:sz'))
                            if sz is not None:
                                half_pt = sz.get(qn('w:val'))
                                if half_pt:
                                    font_size = int(half_pt) / 2.0

                        # 校验
                        if font_name:
                            if "宋体" not in font_name and "SimSun" not in font_name:
                                font_issues.append(
                                    f"脚注「{run_text[:20]}」字体应为宋体，"
                                    f"实际为 {font_name}"
                                )
                        if font_size and abs(font_size - 9.0) > 0.5:
                            font_issues.append(
                                f"脚注「{run_text[:20]}」字号应为小五(9pt)，"
                                f"实际为 {font_size}pt"
                            )

                if scanned > 50:
                    break

            if scanned == 0:
                self.report.add_error("脚注内容为空，跳过脚注校验")
            elif not font_issues:
                self.report.add_pass(
                    "脚注格式",
                    f"全部 {scanned} 条脚注格式符合要求（宋体、小五号）"
                )
            else:
                self.report.add_fail(
                    "脚注格式",
                    f"共 {len(font_issues)} 处脚注格式不合规",
                    context_text="\n".join(font_issues[:15])
                )

        except Exception as e:
            self.report.add_error(f"脚注校验异常: {str(e)}")

    # ─────────────────────────────────────────────────────────
    # Run 级字体解析辅助方法
    # ─────────────────────────────────────────────────────────

    @staticmethod
    def _get_run_east_asian_font(run) -> Optional[str]:
        """获取 run 的东亚字体名（宋体、黑体等）"""
        try:
            rPr = run._element.rPr
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    ea = rFonts.get(qn('w:eastAsia'))
                    if ea:
                        return ea
        except Exception:
            pass
        return None

    @staticmethod
    def _get_run_effective_font(run) -> Optional[str]:
        """获取 run 的有效字体名（优先东亚字体，其次 font.name）"""
        ea = ThesisFormatVerifier._get_run_east_asian_font(run)
        if ea:
            return ea
        return run.font.name

    @staticmethod
    def _get_run_font_size_pt(run) -> Optional[float]:
        """获取 run 的字号（磅）"""
        if run.font.size:
            return run.font.size.pt
        return None

    @staticmethod
    def _is_run_latin_or_digit(run) -> bool:
        """判断 run 的文本是否为纯英文/数字/标点（无中文字符）"""
        text = run.text
        if not text:
            return False
        for ch in text:
            if '一' <= ch <= '鿿' or '㐀' <= ch <= '䶿':
                return False  # 含中文字符
        # 至少包含一个字母或数字才视为英文 run
        return any(c.isalpha() or c.isdigit() for c in text)

    @staticmethod
    def _get_paragraph_line_spacing_multiple(para) -> Optional[float]:
        """
        获取段落行距倍数。
        返回 1.5 表示 1.5 倍行距，None 表示未设定或无法判断。
        """
        pf = para.paragraph_format
        if pf.line_spacing_rule == WD_LINE_SPACING.MULTIPLE and pf.line_spacing:
            return float(pf.line_spacing)
        if pf.line_spacing_rule == WD_LINE_SPACING.EXACTLY and pf.line_spacing:
            # 固定值行距：以 12pt 小四为基准换算倍数
            return pf.line_spacing.pt / 12.0
        return None

    @staticmethod
    def _get_paragraph_indent_cm(para) -> Optional[float]:
        """获取段落首行缩进（厘米），None 表示未设定"""
        indent = para.paragraph_format.first_line_indent
        if indent is None:
            return None
        # python-docx 返回 EMU 值，1 cm = 360000 EMU
        return indent / 360000

    @staticmethod
    def _classify_heading_level(para_text: str) -> Optional[int]:
        """
        判断段落文本是否为标题并返回级别。

        返回:
            1 — 一级标题（如 "1 标题"、"一、标题"）
            2 — 二级标题（如 "1.1 标题"、"（一）标题"）
            3 — 三级标题（如 "1.1.1 标题"、"（1）标题"）
            None — 非标题
        """
        text = para_text.strip()
        if not text:
            return None

        # 三级优先判定（最具体的模式先匹配）
        if ThesisFormatVerifier.RE_HEADING_LEVEL3.match(text):
            return 3

        # 二级：n.n 开头 或 （一）（1）开头
        if re.match(r"^\s*\d+\.\d+\s", text):
            return 2
        if re.match(r"^\s*（[一二三四五六七八九十百]+）", text):
            return 2
        if re.match(r"^\s*\([一二三四五六七八九十百]+\)", text):
            return 2

        # 一级：单数字+分隔符 或 中文序号+顿号
        if re.match(r"^\s*\d+\s*[、.,，]", text):
            # 排除 n.n 模式（已归为二级）
            if not re.match(r"^\s*\d+\.\d+", text):
                return 1
        if re.match(r"^\s*[一二三四五六七八九十百]+[、．.]", text):
            return 1

        return None

    # ─────────────────────────────────────────────────────────
    # 校验项：全局正文字体与行距
    # ─────────────────────────────────────────────────────────

    def _check_fonts_and_spacing(self):
        """
        遍历正文段落，校验：
        1. 行距为 1.5 倍
        2. 中文 run 字体为宋体，字号为小四（12pt）
        3. 英文/数字 run 字体为 Times New Roman，字号为小四（12pt）

        跳过标题段落（由 _check_heading_styles 负责）。
        """
        if not self.doc.paragraphs:
            self.report.add_error("文档无段落，跳过正文字体与行距校验")
            return

        body_started = False
        line_spacing_issues: List[str] = []
        font_issues: List[str] = []
        scanned_count = 0

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 粗略定位正文起点：跳过封面/声明/目录/摘要等前导部分
            # 遇到第一个 "正文" 或 "第一章" 或 "绪论" 等才开始
            if not body_started:
                if any(kw in text for kw in ["第一章", "第1章", "绪论", "引言",
                                              "第 一 章", "1 绪论", "1 绪"]):
                    body_started = True
                else:
                    continue

            # 跳过标题段落（由 _check_heading_styles 专门校验）
            if self._classify_heading_level(text) is not None:
                continue

            # 跳过参考文献及之后的内容
            if text.startswith("参考文献") or text.startswith("致谢") or text.startswith("致 谢"):
                break

            scanned_count += 1
            if scanned_count > 500:  # 安全上限，避免超长文档卡死
                break

            # ── 行距校验 ──
            spacing = self._get_paragraph_line_spacing_multiple(para)
            if spacing is None:
                line_spacing_issues.append(f"段落「{text[:30]}…」行距未设定")
            elif abs(spacing - 1.5) > 0.05:
                line_spacing_issues.append(
                    f"段落「{text[:30]}…」行距为 {spacing:.2f}，应为 1.5"
                )

            # ── 字体与字号校验（Run 级深度解析）──
            for run in para.runs:
                run_text = run.text
                if not run_text or not run_text.strip():
                    continue

                font_size = self._get_run_font_size_pt(run)

                if self._is_run_latin_or_digit(run):
                    # 英文/数字 run：Times New Roman，小四(12pt)
                    font_name = self._get_run_effective_font(run)
                    if font_name and "Times New Roman" not in font_name:
                        font_issues.append(
                            f"「{run_text[:20]}」英文应为 Times New Roman，"
                            f"实际为 {font_name}"
                        )
                    if font_size and abs(font_size - 12.0) > self.FONT_SIZE_TOLERANCE:
                        font_issues.append(
                            f"「{run_text[:20]}」字号应为小四(12pt)，"
                            f"实际为 {font_size}pt"
                        )
                else:
                    # 中文 run：宋体，小四(12pt)
                    font_name = self._get_run_effective_font(run)
                    if font_name and "宋体" not in font_name and "SimSun" not in font_name:
                        font_issues.append(
                            f"「{run_text[:20]}」中文应为宋体，"
                            f"实际为 {font_name}"
                        )
                    if font_size and abs(font_size - 12.0) > self.FONT_SIZE_TOLERANCE:
                        font_issues.append(
                            f"「{run_text[:20]}」字号应为小四(12pt)，"
                            f"实际为 {font_size}pt"
                        )

        # ── 汇总行距结果 ──
        if not line_spacing_issues:
            self.report.add_pass(
                "正文行距",
                f"已检查 {scanned_count} 个正文段落，行距均为 1.5 倍"
            )
        else:
            self.report.add_fail(
                "正文行距",
                f"共 {len(line_spacing_issues)} 个段落行距不合规",
                context_text="\n".join(line_spacing_issues[:15])
            )

        # ── 汇总字体结果 ──
        if not font_issues:
            self.report.add_pass(
                "正文字体",
                f"正文中文(宋体/小四)、英文(Times New Roman/小四) 均符合规范"
            )
        else:
            self.report.add_fail(
                "正文字体",
                f"共 {len(font_issues)} 处字体或字号不合规",
                context_text="\n".join(font_issues[:15])
            )

    # ─────────────────────────────────────────────────────────
    # 校验项：多级标题样式与缩进
    # ─────────────────────────────────────────────────────────

    def _check_heading_styles(self):
        """
        校验多级标题：
        - 一级标题：黑体，小三号(15pt)，1.5 倍行距
        - 二级标题：黑体，四号(14pt)，1.5 倍行距
        - 三级标题：黑体，小四号(12pt)，1.5 倍行距
        - 缩进规则：阿拉伯数字序号顶格，中文序号缩进两个中文字符
        """
        if not self.doc.paragraphs:
            self.report.add_error("文档无段落，跳过标题样式校验")
            return

        # 按级别分组收集问题
        level_config = {
            1: {"name": "一级标题", "font": "黑体", "size_pt": 15.0,
                "size_name": "小三", "issues": [], "count": 0, "contexts": []},
            2: {"name": "二级标题", "font": "黑体", "size_pt": 14.0,
                "size_name": "四号", "issues": [], "count": 0, "contexts": []},
            3: {"name": "三级标题", "font": "黑体", "size_pt": 12.0,
                "size_name": "小四", "issues": [], "count": 0, "contexts": []},
        }

        body_started = False

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 定位正文起点
            if not body_started:
                if any(kw in text for kw in ["第一章", "第1章", "绪论", "引言",
                                              "第 一 章", "1 绪论"]):
                    body_started = True
                else:
                    continue

            # 遇到参考文献/致谢则停止
            if text.startswith("参考文献") or text.startswith("致谢") or text.startswith("致 谢"):
                break

            level = self._classify_heading_level(text)
            if level is None:
                continue

            cfg = level_config[level]
            cfg["count"] += 1
            issues = []

            # ── 字体校验（黑体）──
            font_name = None
            font_size = None
            for run in para.runs:
                if not run.text.strip():
                    continue
                if font_name is None:
                    font_name = self._get_run_effective_font(run)
                if font_size is None:
                    font_size = self._get_run_font_size_pt(run)

            if font_name and "黑体" not in font_name and "SimHei" not in font_name:
                issues.append(f"字体应为黑体，实际为 {font_name}")

            # ── 字号校验 ──
            if font_size and abs(font_size - cfg["size_pt"]) > self.FONT_SIZE_TOLERANCE:
                issues.append(
                    f"字号应为{cfg['size_name']}({cfg['size_pt']}pt)，"
                    f"实际为 {font_size}pt"
                )

            # ── 行距校验（1.5 倍）──
            spacing = self._get_paragraph_line_spacing_multiple(para)
            if spacing is None:
                issues.append("行距未设定，应为 1.5 倍")
            elif abs(spacing - 1.5) > 0.05:
                issues.append(f"行距应为 1.5 倍，实际为 {spacing:.2f}")

            # ── 缩进校验（双轨制）──
            indent_cm = self._get_paragraph_indent_cm(para)

            if self.RE_HEADING_ARABIC.match(text):
                # 阿拉伯数字格式 → 必须顶格（缩进应为 0 或 None）
                if indent_cm is not None and abs(indent_cm) > 0.05:
                    issues.append(
                        f"阿拉伯数字序号标题应顶格，"
                        f"实际缩进 {indent_cm:.2f}cm"
                    )
            elif self.RE_HEADING_CHINESE.match(text):
                # 中文序号格式 → 必须缩进两个中文字符
                # 宽度随字号变化：2 × font_size_pt × 0.0353 cm/pt
                # 容差 ±0.2cm
                expected_indent = 2 * cfg["size_pt"] * 0.0353
                indent_tolerance = 0.2
                if indent_cm is None:
                    issues.append("中文序号标题应缩进两个中文字符，实际未设定缩进")
                elif abs(indent_cm - expected_indent) > indent_tolerance:
                    issues.append(
                        f"中文序号标题应缩进两个中文字符(≈0.74cm)，"
                        f"实际缩进 {indent_cm:.2f}cm"
                    )

            if issues:
                cfg["issues"].append(f"「{text[:40]}」" + "；".join(issues))
                cfg["contexts"].append(f"[{text[:40]}] 字体:{font_name}, 字号:{font_size}pt")

        # ── 汇总各级标题结果 ──
        for level, cfg in level_config.items():
            if cfg["count"] == 0:
                self.report.add_error(f"未检测到任何{cfg['name']}，跳过校验")
                continue
            if not cfg["issues"]:
                self.report.add_pass(
                    f"{cfg['name']}样式",
                    f"全部 {cfg['count']} 个{cfg['name']}格式符合要求"
                    f"（{cfg['font']}/{cfg['size_name']}/1.5倍行距/缩进正确）"
                )
            else:
                self.report.add_fail(
                    f"{cfg['name']}样式",
                    f"共 {len(cfg['issues'])} 个{cfg['name']}格式不合规"
                    f"（共检测到 {cfg['count']} 个）",
                    context_text="\n".join(cfg["contexts"][:10])
                )

    # ─────────────────────────────────────────────────────────
    # 辅助：定位正文在段落列表中的起止索引
    # ─────────────────────────────────────────────────────────

    def _find_body_range(self) -> Tuple[int, int]:
        """
        返回正文段落在 self.doc.paragraphs 中的 [start, end) 索引区间。
        start: 第一个正文段落（含 "第一章" 等关键词的段落）
        end:   "参考文献" 标题段落索引（不含），若未找到则到文档末尾。
        """
        paras = self.doc.paragraphs
        start = None
        end = len(paras)

        for i, p in enumerate(paras):
            text = p.text.strip()
            if not text:
                continue
            if start is None:
                if any(kw in text for kw in ["第一章", "第1章", "绪论", "引言",
                                              "第 一 章", "1 绪论"]):
                    start = i
            else:
                if text.startswith("参考文献"):
                    end = i
                    break

        if start is None:
            return (0, 0)
        return (start, end)

    @staticmethod
    def _is_caption_paragraph(para, keyword: str) -> bool:
        """判断段落是否是题注段落（包含表序/图序关键字，如 '表1' 或 '图1'）"""
        text = para.text.strip()
        if not text:
            return False
        # 匹配 "表1" "表 1" "表1-1" "表 1.2" 以及 "图1" 等变体
        pattern = rf"^{re.escape(keyword)}\s*\d"
        return bool(re.match(pattern, text))

    @staticmethod
    def _check_caption_format(para, expected_font_keywords: List[str],
                               expected_size_pt: float,
                               font_size_tolerance: float = 0.5) -> List[str]:
        """
        校验题注段落的格式：居中、字体、字号。
        返回问题列表（空列表表示合规）。
        """
        issues = []
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            align_name = {
                WD_ALIGN_PARAGRAPH.LEFT: "左对齐",
                WD_ALIGN_PARAGRAPH.RIGHT: "右对齐",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "两端对齐",
            }.get(para.alignment, "未设定")
            issues.append(f"应居中，实际为{align_name}")

        font_name = None
        font_size = None
        for run in para.runs:
            if not run.text.strip():
                continue
            if font_name is None:
                ea = ThesisFormatVerifier._get_run_east_asian_font(run)
                font_name = ea if ea else run.font.name
            if font_size is None and run.font.size:
                font_size = run.font.size.pt

        if font_name:
            if not any(kw in font_name for kw in expected_font_keywords):
                issues.append(f"字体应为{'/'.join(expected_font_keywords)}，实际为 {font_name}")
        if font_size and abs(font_size - expected_size_pt) > font_size_tolerance:
            issues.append(f"字号应为五号({expected_size_pt}pt)，实际为 {font_size}pt")

        return issues

    # ─────────────────────────────────────────────────────────
    # 校验项：图表题注位置与格式
    # ─────────────────────────────────────────────────────────

    def _check_tables_and_figures(self):
        """
        校验表格和插图的题注规范：

        表格：表序与表名居中放在表格上方，宋体/TNR，五号(10.5pt)，
              表序与表名之间空一个中文字符。
        插图：图序与图名居中放在图的下方，宋体，五号(10.5pt)，
              图序与图名之间空一个中文字符。
        """
        body_start, body_end = self._find_body_range()

        # ── 表格题注校验 ──
        # 策略：遍历 doc.element.body 的 block-level 元素，
        # 遇到 <w:tbl> 时检查紧邻的前一个 <w:p> 是否为合规的表题注。
        table_issues: List[str] = []
        from docx.oxml.ns import qn as _qn
        body_element = self.doc.element.body
        block_elements = list(body_element)

        prev_para_text = None
        prev_para_element = None
        table_seq = 0

        for elem in block_elements:
            tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag

            if tag == 'tbl':
                table_seq += 1
                # 检查前一个段落是否为表题注
                if prev_para_text and self._is_caption_paragraph(prev_para_text, "表"):
                    # 校验格式
                    if prev_para_element is not None:
                        # 构造一个临时段落对象来校验格式
                        para_obj = None
                        for p in self.doc.paragraphs:
                            if p._element is prev_para_element:
                                para_obj = p
                                break
                        if para_obj:
                            fmt_issues = self._check_caption_format(
                                para_obj,
                                ["宋体", "SimSun", "Times New Roman"],
                                10.5
                            )
                            # 检查表序与表名之间的空格
                            caption_text = para_obj.text.strip()
                            # 去掉 "表1" 部分，看后面是否有空格
                            m = re.match(r"^(表\s*\d[\d\-\.]*)\s*(.*)", caption_text)
                            if m:
                                sep = m.group(2)
                                if sep and not sep.startswith("　") and not sep.startswith(" "):
                                    fmt_issues.append("表序与表名之间应空一个中文字符")

                            if fmt_issues:
                                table_issues.append(
                                    f"表{table_seq}「{prev_para_text.text.strip()[:30]}」"
                                    + "；".join(fmt_issues)
                                )
                        else:
                            table_issues.append(f"表{table_seq}：无法定位题注段落进行格式校验")
                    else:
                        table_issues.append(f"表{table_seq}：上方无段落元素")
                elif prev_para_text is None:
                    table_issues.append(f"表{table_seq}：表格位于文档开头，无上方段落")
                else:
                    table_issues.append(
                        f"表{table_seq}：上方段落「{prev_para_text.text.strip()[:30]}」"
                        f"不像是表题注（应以'表'开头）"
                    )

            # 更新"前一个段落"引用
            if tag == 'p':
                prev_para_element = elem
                # 提取纯文本
                texts = []
                for t_elem in elem.iter(_qn('w:t')):
                    if t_elem.text:
                        texts.append(t_elem.text)
                # 创建简易对象以存储文本
                class _SimplePara:
                    def __init__(self, text, element):
                        self.text = text
                        self._element = element
                        self.alignment = None
                        self.runs = []
                        self.paragraph_format = None
                    def strip(self):
                        return self.text.strip()
                prev_para_text = _SimplePara("".join(texts), elem)

        if table_seq == 0:
            self.report.add_error("文档中未检测到表格，跳过表格题注校验")
        elif not table_issues:
            self.report.add_pass(
                "表格题注",
                f"全部 {table_seq} 个表格的题注格式符合规范"
                f"（居中、宋体/TNR、五号、位于表格上方）"
            )
        else:
            self.report.add_fail(
                "表格题注",
                f"共 {len(table_issues)} 个表格题注不合规（共 {table_seq} 个表格）",
                context_text="\n".join(table_issues[:10])
            )

        # ── 插图题注校验 ──
        # python-docx 没有直接的 "figure" 遍历接口，
        # 通过 inline shapes 或者段落中含 "图{n}" 关键字来检测。
        # 策略：扫描正文段落，找到 "图{n}" 开头的段落，
        # 然后检查其前一个非空段落是否包含图片（inline shape）。
        figure_issues: List[str] = []
        figure_count = 0

        for i in range(body_start + 1, body_end):
            para = self.doc.paragraphs[i]
            text = para.text.strip()
            if not text:
                continue

            if self._is_caption_paragraph(para, "图"):
                figure_count += 1
                # 图题注应在图的下方，即当前段落是题注，
                # 图片应在前一个段落中（或前几个段落中）
                # 检查前面的段落是否含图片
                has_image_above = False
                if i - 1 >= body_start:
                    prev = self.doc.paragraphs[i - 1]
                    # 检查紧邻的前一个段落是否含 inline shape（图片）
                    for run in prev.runs:
                        if run._element.findall(_qn('w:drawing')):
                            has_image_above = True
                            break
                        if run._element.findall('.//' + _qn('a:blip')):
                            has_image_above = True
                            break

                if not has_image_above:
                    figure_issues.append(
                        f"图{figure_count}「{text[:30]}」：题注上方未检测到图片"
                    )

                # 校验题注格式
                fmt_issues = self._check_caption_format(
                    para, ["宋体", "SimSun"], 10.5
                )
                # 检查图序与图名之间的空格
                m = re.match(r"^(图\s*\d[\d\-\.]*)\s*(.*)", text)
                if m:
                    sep = m.group(2)
                    if sep and not sep.startswith("　") and not sep.startswith(" "):
                        fmt_issues.append("图序与图名之间应空一个中文字符")

                if fmt_issues:
                    figure_issues.append(
                        f"图{figure_count}「{text[:30]}」" + "；".join(fmt_issues)
                    )

        if figure_count == 0:
            self.report.add_error("文档中未检测到图题注（'图n'开头的段落），跳过插图校验")
        elif not figure_issues:
            self.report.add_pass(
                "插图题注",
                f"全部 {figure_count} 个插图的题注格式符合规范"
                f"（居中、宋体、五号、位于图片下方）"
            )
        else:
            self.report.add_fail(
                "插图题注",
                f"共 {len(figure_issues)} 个插图题注不合规（共 {figure_count} 个图题注）",
                context_text="\n".join(figure_issues[:10])
            )

    # ─────────────────────────────────────────────────────────
    # 校验项：参考文献比例指标（核心算法）
    # ─────────────────────────────────────────────────────────

    def _check_references_metrics(self):
        """
        精确定位参考文献区域，逐条解析文献，计算比例指标：

        a) 总数不少于 10 篇
        b) 近 3 年文献不少于 1/3
        c) 外文文献不少于 1/5
        d) 学位论文不多于 1/5
        同时校验标题格式（黑体/五号）、行距（1.5 倍）、悬挂缩进。
        """
        paras = self.doc.paragraphs

        # ── 定位 "参考文献" 标题段落 ──
        ref_title_idx = None
        for i, p in enumerate(paras):
            text = p.text.strip()
            if text == "参考文献" or text == "参考文献 ":
                ref_title_idx = i
                break

        if ref_title_idx is None:
            self.report.add_error("未找到独立的「参考文献」标题段落，跳过参考文献校验")
            return

        # ── 校验标题格式：黑体、五号(10.5pt) ──
        ref_title_para = paras[ref_title_idx]
        title_font = None
        title_size = None
        for run in ref_title_para.runs:
            if not run.text.strip():
                continue
            if title_font is None:
                ea = self._get_run_east_asian_font(run)
                title_font = ea if ea else run.font.name
            if title_size is None and run.font.size:
                title_size = run.font.size.pt

        title_issues = []
        if title_font and "黑体" not in title_font and "SimHei" not in title_font:
            title_issues.append(f"标题字体应为黑体，实际为 {title_font}")
        if title_size and abs(title_size - 10.5) > 0.5:
            title_issues.append(f"标题字号应为五号(10.5pt)，实际为 {title_size}pt")

        if title_issues:
            self.report.add_fail(
                "参考文献标题",
                "；".join(title_issues)
            )
        else:
            self.report.add_pass(
                "参考文献标题",
                "「参考文献」标题格式符合要求（黑体、五号）"
            )

        # ── 收集参考文献条目 ──
        # 从标题下一个段落开始，到遇到下一个主标题或文档末尾
        ref_entries: List[str] = []           # 纯文本
        ref_entry_paras: List = []            # 段落对象
        for i in range(ref_title_idx + 1, len(paras)):
            p = paras[i]
            text = p.text.strip()
            # 遇到新的主标题（致谢、附录等）则停止
            if text and (text.startswith("致谢") or text.startswith("致 谢")
                         or text.startswith("附录") or text.startswith("附 录")):
                break
            # 跳过空行
            if not text:
                continue
            ref_entries.append(text)
            ref_entry_paras.append(p)

        if not ref_entries:
            self.report.add_error("参考文献区域为空，跳过比例校验")
            return

        total = len(ref_entries)

        # ── 条目格式校验：行距 1.5 倍 + 悬挂缩进 ──
        entry_format_issues: List[str] = []
        for idx, (text, para) in enumerate(zip(ref_entries, ref_entry_paras)):
            # 行距
            spacing = self._get_paragraph_line_spacing_multiple(para)
            if spacing is not None and abs(spacing - 1.5) > 0.05:
                entry_format_issues.append(
                    f"[{idx+1}] 行距为 {spacing:.2f}，应为 1.5"
                )
            # 悬挂缩进（first_line_indent < 0 表示悬挂缩进）
            indent = para.paragraph_format.first_line_indent
            if indent is not None and indent >= 0:
                snippet = text[:40] + ("..." if len(text) > 40 else "")
                entry_format_issues.append(
                    f"[{idx+1}] 「{snippet}」未设置悬挂缩进"
                )

        if not entry_format_issues:
            self.report.add_pass(
                "参考文献条目格式",
                f"全部 {total} 条文献行距(1.5倍)与悬挂缩进均符合要求"
            )
        else:
            self.report.add_fail(
                "参考文献条目格式",
                f"共 {len(entry_format_issues)} 条文献格式不合规",
                context_text="\n".join(entry_format_issues[:15])
            )

        # ── a) 总数校验 ──
        min_total = 10
        if total >= min_total:
            self.report.add_pass(
                "参考文献总数",
                f"共 {total} 篇，满足 ≥{min_total} 篇要求"
            )
        else:
            self.report.add_fail(
                "参考文献总数",
                f"参考文献总数不足 {min_total} 篇，当前为 {total} 篇"
            )

        # ── b) 近 3 年文献比例 ──
        current_year = datetime.now().year
        recent_year_start = current_year - 2  # 例如 2026 → 2024
        year_pattern = re.compile(r"(19\d{2}|20\d{2})")
        recent_count = 0
        year_found_count = 0
        old_refs: List[str] = []

        for idx, text in enumerate(ref_entries):
            years = year_pattern.findall(text)
            if years:
                year_found_count += 1
                latest = max(int(y) for y in years)
                if recent_year_start <= latest <= current_year:
                    recent_count += 1
                else:
                    old_refs.append(f"[{idx+1}] {text[:60]}")

        min_recent_ratio = 1 / 3
        if year_found_count > 0:
            recent_ratio = recent_count / year_found_count
            if recent_ratio >= min_recent_ratio:
                self.report.add_pass(
                    "参考文献年份",
                    f"近3年(≥{recent_year_start})文献 {recent_count}/{year_found_count} 篇，"
                    f"占比 {recent_ratio:.1%}，满足 ≥{min_recent_ratio:.0%} 要求"
                )
            else:
                self.report.add_fail(
                    "参考文献年份",
                    f"近3年(≥{recent_year_start})文献 {recent_count}/{year_found_count} 篇，"
                    f"占比 {recent_ratio:.1%}，不满足 ≥{min_recent_ratio:.0%} 要求"
                    f"（需至少 {int(total * min_recent_ratio) + 1} 篇近3年文献）",
                    context_text="\n".join(old_refs[:10]) if old_refs else None
                )
        else:
            self.report.add_error("参考文献中未能提取到有效年份信息，跳过年份比例校验")

        # ── c) 外文文献比例 ──
        min_foreign_ratio = 1 / 5
        foreign_count = 0
        non_foreign_refs: List[str] = []

        for idx, text in enumerate(ref_entries):
            alpha_chars = [c for c in text if c.isalpha()]
            if alpha_chars:
                english_chars = [c for c in alpha_chars if ord(c) < 128]
                # 英文字符占比超过 50% 视为外文文献
                if len(english_chars) / len(alpha_chars) > 0.5:
                    foreign_count += 1
                else:
                    non_foreign_refs.append(f"[{idx+1}] {text[:60]}")
            else:
                non_foreign_refs.append(f"[{idx+1}] {text[:60]}")

        foreign_ratio = foreign_count / total
        if foreign_ratio >= min_foreign_ratio:
            self.report.add_pass(
                "外文文献比例",
                f"外文文献 {foreign_count}/{total} 篇，占比 {foreign_ratio:.1%}，"
                f"满足 ≥{min_foreign_ratio:.0%} 要求"
            )
        else:
            self.report.add_fail(
                "外文文献比例",
                f"外文文献 {foreign_count}/{total} 篇，占比 {foreign_ratio:.1%}，"
                f"不满足 ≥{min_foreign_ratio:.0%} 要求"
                f"（需至少 {int(total * min_foreign_ratio) + 1} 篇外文文献）",
                context_text=(
                    f"以下为非外文文献（共 {len(non_foreign_refs)} 篇）：\n"
                    + "\n".join(non_foreign_refs[:10])
                ) if non_foreign_refs else None
            )

        # ── d) 学位论文比例 ──
        max_thesis_ratio = 1 / 5
        thesis_tag = "[D]"
        thesis_refs: List[str] = []
        for idx, text in enumerate(ref_entries):
            if thesis_tag in text:
                thesis_refs.append(f"[{idx+1}] {text[:60]}")

        thesis_count = len(thesis_refs)
        thesis_ratio = thesis_count / total
        if thesis_ratio <= max_thesis_ratio:
            self.report.add_pass(
                "学位论文比例",
                f"学位论文 {thesis_count}/{total} 篇，占比 {thesis_ratio:.1%}，"
                f"满足 ≤{max_thesis_ratio:.0%} 要求"
            )
        else:
            self.report.add_fail(
                "学位论文比例",
                f"学位论文 {thesis_count}/{total} 篇，占比 {thesis_ratio:.1%}，"
                f"不满足 ≤{max_thesis_ratio:.0%} 要求"
                f"（当前超出 {thesis_count - int(total * max_thesis_ratio)} 篇）",
                context_text="\n".join(thesis_refs) if thesis_refs else None
            )


# ─────────────────────────────────────────────────────────────
# CLI 入口
# ─────────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python Thesis_verify_tool.py <path_to_docx>")
        print("\n--- 校验项 ---")
        print("  1. 全局页边距（上/下 2.54cm，左/右 3.17cm）")
        print("  2. 文档结构顺序（封面→声明→目录→摘要→正文→参考文献→致谢→附录）")
        print("  3. 封面格式（主标题宋体/一号/加粗，日期阿拉伯数字）")
        print("  4. 中文题目与摘要（楷体_GB2312、二号/小四）")
        print("  5. 英文题目与摘要（Times New Roman、二号/加粗）")
        print("  6. 正文字体与行距（宋体/TNR、小四、1.5倍行距）")
        print("  7. 多级标题样式与缩进（黑体、字号、双轨制缩进）")
        print("  8. 图表题注位置与格式（表上方/图下方、居中、五号）")
        print("  9. 参考文献比例指标（总数/近3年/外文/学位论文）")
        print("  10. 脚注格式（宋体、小五号）")
        sys.exit(1)

    file_path = sys.argv[1]

    verifier = ThesisFormatVerifier(file_path)
    result = verifier.run_all()

    # 控制台输出
    summary = result["summary"]
    print(f"\n{'='*50}")
    print(f"  毕业论文格式校验报告")
    print(f"{'='*50}")
    print(f"  总检查项: {summary['total_checks']}")
    print(f"  通过: {summary['passed']}")
    print(f"  失败: {summary['failed']}")
    print(f"  异常: {summary['errors']}")
    print(f"{'='*50}\n")

    if result["passed_items"]:
        print("[通过]")
        for item in result["passed_items"]:
            print(f"  ✓ {item['name']}: {item['message']}")

    if result["failed_items"]:
        print("\n[失败]")
        for item in result["failed_items"]:
            print(f"  ✗ {item['name']}: {item['message']}")
            if "context_text" in item:
                for line in item["context_text"].split("\n"):
                    print(f"      {line}")

    if result["errors"]:
        print("\n[异常]")
        for err in result["errors"]:
            print(f"  ! {err}")
